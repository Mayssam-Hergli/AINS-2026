"""Document extractors for the File Parser Service.

Each extractor turns one supported file type into a uniform :class:`RawExtraction`
(text blocks + tables + metadata) which the :class:`~services.file_parser.normalizer.Normalizer`
then converts into typed signals.

The underlying libraries (``pdfplumber``, ``python-docx``, ``pandas``) are all
synchronous and CPU/IO-bound, so every extractor exposes an ``async`` ``extract``
method that offloads the heavy work to a thread via :func:`asyncio.to_thread`,
keeping the event loop free for concurrent document processing.
"""

from __future__ import annotations

import asyncio
from abc import ABC, abstractmethod
from io import BytesIO
from pathlib import PurePath

from pydantic import BaseModel, Field


class TextBlock(BaseModel):
    """A contiguous run of prose with a locator describing where it came from."""

    text: str
    locator: str = ""


class Table(BaseModel):
    """A tabular region, rows-first. The first row is treated as a header if present."""

    rows: list[list[str]]
    locator: str = ""


class RawExtraction(BaseModel):
    """Library-agnostic representation of one parsed file."""

    document_id: str
    filename: str
    content_type: str
    text_blocks: list[TextBlock] = Field(default_factory=list)
    tables: list[Table] = Field(default_factory=list)
    metadata: dict[str, str] = Field(default_factory=dict)


class UnsupportedDocumentError(ValueError):
    """Raised when no registered extractor can handle a file."""


class BaseExtractor(ABC):
    """Common contract for all extractors."""

    #: File extensions (lower-case, no dot) this extractor handles.
    extensions: tuple[str, ...] = ()
    #: MIME types this extractor handles.
    content_types: tuple[str, ...] = ()

    async def extract(
        self, *, document_id: str, filename: str, content_type: str, data: bytes
    ) -> RawExtraction:
        """Parse ``data`` off the event loop and return a :class:`RawExtraction`."""
        return await asyncio.to_thread(
            self._extract_sync,
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            data=data,
        )

    @abstractmethod
    def _extract_sync(
        self, *, document_id: str, filename: str, content_type: str, data: bytes
    ) -> RawExtraction:
        """Blocking extraction; runs inside a worker thread."""


class PDFExtractor(BaseExtractor):
    """Extracts text and tables from PDFs using ``pdfplumber``."""

    extensions = ("pdf",)
    content_types = ("application/pdf",)

    def _extract_sync(
        self, *, document_id: str, filename: str, content_type: str, data: bytes
    ) -> RawExtraction:
        try:
            import pdfplumber
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise RuntimeError("pdfplumber is required to parse PDF files") from exc

        text_blocks: list[TextBlock] = []
        tables: list[Table] = []
        metadata: dict[str, str] = {}

        with pdfplumber.open(BytesIO(data)) as pdf:
            doc_meta = pdf.metadata or {}
            metadata = {str(k): str(v) for k, v in doc_meta.items() if v is not None}
            metadata["page_count"] = str(len(pdf.pages))

            for index, page in enumerate(pdf.pages, start=1):
                locator = f"page:{index}"
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_blocks.append(TextBlock(text=page_text, locator=locator))
                for raw_table in page.extract_tables() or []:
                    rows = [
                        [(cell or "").strip() for cell in row]
                        for row in raw_table
                        if any(cell for cell in row)
                    ]
                    if rows:
                        tables.append(Table(rows=rows, locator=locator))

        return RawExtraction(
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            text_blocks=text_blocks,
            tables=tables,
            metadata=metadata,
        )


class DocxExtractor(BaseExtractor):
    """Extracts paragraphs and tables from Word documents using ``python-docx``."""

    extensions = ("docx",)
    content_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def _extract_sync(
        self, *, document_id: str, filename: str, content_type: str, data: bytes
    ) -> RawExtraction:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise RuntimeError("python-docx is required to parse DOCX files") from exc

        document = Document(BytesIO(data))
        text_blocks: list[TextBlock] = []
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if text:
                text_blocks.append(TextBlock(text=text, locator=f"para:{index}"))

        tables: list[Table] = []
        for t_index, table in enumerate(document.tables):
            rows = [
                [cell.text.strip() for cell in row.cells] for row in table.rows
            ]
            rows = [row for row in rows if any(row)]
            if rows:
                tables.append(Table(rows=rows, locator=f"table:{t_index}"))

        core = document.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "paragraph_count": str(len(document.paragraphs)),
        }

        return RawExtraction(
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            text_blocks=text_blocks,
            tables=tables,
            metadata={k: v for k, v in metadata.items() if v},
        )


class TabularExtractor(BaseExtractor):
    """Extracts sheets/rows from CSV and Excel files using ``pandas``."""

    extensions = ("csv", "xlsx", "xls")
    content_types = (
        "text/csv",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    def _extract_sync(
        self, *, document_id: str, filename: str, content_type: str, data: bytes
    ) -> RawExtraction:
        try:
            import pandas as pd
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise RuntimeError("pandas is required to parse tabular files") from exc

        extension = PurePath(filename).suffix.lower().lstrip(".")
        sheets: dict[str, "pd.DataFrame"]
        if extension == "csv" or content_type == "text/csv":
            sheets = {"data": pd.read_csv(BytesIO(data))}
        else:
            sheets = pd.read_excel(BytesIO(data), sheet_name=None)  # all sheets

        tables: list[Table] = []
        total_rows = 0
        for sheet_name, frame in sheets.items():
            frame = frame.fillna("")
            header = [str(col) for col in frame.columns]
            rows = [header] + [
                [str(cell) for cell in record] for record in frame.itertuples(index=False)
            ]
            total_rows += len(frame)
            tables.append(Table(rows=rows, locator=f"sheet:{sheet_name}"))

        return RawExtraction(
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            tables=tables,
            metadata={"sheet_count": str(len(sheets)), "row_count": str(total_rows)},
        )


class ExtractorRegistry:
    """Resolves the right extractor for a file by extension or MIME type."""

    def __init__(self, extractors: list[BaseExtractor] | None = None) -> None:
        self._extractors = extractors or [
            PDFExtractor(),
            DocxExtractor(),
            TabularExtractor(),
        ]
        self._by_extension: dict[str, BaseExtractor] = {}
        self._by_content_type: dict[str, BaseExtractor] = {}
        for extractor in self._extractors:
            for ext in extractor.extensions:
                self._by_extension[ext] = extractor
            for ctype in extractor.content_types:
                self._by_content_type[ctype] = extractor

    def resolve(self, *, filename: str, content_type: str | None) -> BaseExtractor:
        if content_type and content_type in self._by_content_type:
            return self._by_content_type[content_type]
        extension = PurePath(filename).suffix.lower().lstrip(".")
        if extension in self._by_extension:
            return self._by_extension[extension]
        raise UnsupportedDocumentError(
            f"No extractor for filename={filename!r} content_type={content_type!r}"
        )

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return tuple(sorted(self._by_extension))
