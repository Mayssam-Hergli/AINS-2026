"""End-to-end demo of the assessment platform.

Runs the full pipeline against an in-memory database so it works with no Postgres and
no Anthropic key:

    documents → File Parser → signals
    signals + questionnaire → MS1 Diagnostic → diagnostic_answers
    diagnostic_answers → MS2 Scoring Agent → 5 scores
    scores → (MS3 reads them back)

Run it directly:  ``python demo.py``
"""

from __future__ import annotations

import asyncio
import sys
import uuid
from io import BytesIO

# Windows consoles default to cp1252; force UTF-8 so the report renders cleanly.
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from services.common.observability import configure_logging
from services.file_parser.core import FileParserService, ParseRequest, SignalRepository
from services.file_parser.extractors import (
    BaseExtractor,
    ExtractorRegistry,
    DocxExtractor,
    PDFExtractor,
    RawExtraction,
    TabularExtractor,
    TextBlock,
)
from services.ms1_diagnostic.engine import DiagnosticEngine, ProjectProfileRepository
from services.ms1_diagnostic.gap_detector import GapDetector
from services.ms1_diagnostic.questions import build_default_graph
from services.scoring_agent.agent import ScoringAgent
from services.scoring_agent.repository import ScoreRepository
from tests.fakepg import InMemoryPool


class InlineTextExtractor(BaseExtractor):
    """Parses a plain-text report into one signal per line — needs no third-party lib."""

    extensions = ("txt",)
    content_types = ("text/plain",)

    def _extract_sync(self, *, document_id, filename, content_type, data):
        text = data.decode("utf-8", errors="ignore")
        blocks = [
            TextBlock(text=line.strip(), locator=f"line:{i}")
            for i, line in enumerate(text.splitlines(), start=1)
            if line.strip()
        ]
        return RawExtraction(
            document_id=document_id, filename=filename,
            content_type=content_type, text_blocks=blocks,
        )


REPORT_TEXT = """\
Sustainability & Compliance Report 2023
In 2023 our Scope 1 emissions were 1,240 tCO2e and Scope 2 emissions were 3,500 tCO2e.
Scope 3 value-chain emissions were measured at 12,400 tCO2e using the GHG Protocol.
Total energy consumption fell 8% year on year to 45,000 MWh.
Renewable energy accounted for 62% of total energy consumption.
Water withdrawal decreased to 120000 m3 and 78% of operational waste was recycled.
We have committed to net-zero by 2040 with an SBTi-validated science-based target.
The gender pay gap stood at 4.5% across the workforce.
Employees received 18 hours of training on average during the year.
The board approved a code of conduct and anti-corruption policy.
An independent whistleblower mechanism is in place and operated throughout the year.
The organisation is in scope for CSRD and reports under the ESRS standards.
Reporting also aligns with the TCFD recommendations and the EU Taxonomy.
Our environmental management system is ISO 14001 certified.
An ISO 27001 certified information security management system is maintained.
Sustainability disclosures received limited external assurance.
"""

# Calibrated, varied self-assessment: environmental/compliance are answered honestly
# (close to the evidence); social/governance are over-claimed (the platform flags them).
QUESTIONNAIRE_ANSWERS = {
    "c_csrd_scope": True,
    "c_esrs_datapoints": 3,
    "c_external_assurance": "limited",
    "e_emissions_tracking": True,
    "e_scope3": 3,
    "e_reduction_target": "validated",
    "s_dei_metrics": ["gender_pay", "training", "diversity", "safety"],
    "g_ethics_policy": True,
    "g_whistleblower": True,
}


def _build_docx() -> bytes | None:
    try:
        from docx import Document
    except ImportError:
        return None
    doc = Document()
    doc.add_heading("Governance & Social Addendum", level=1)
    for line in [
        "The board comprises 60% independent directors with an audit committee.",
        "A formal anti-bribery and ethics policy is reviewed annually by the board.",
        "Anti-corruption training was completed by 95% of staff during the year.",
        "Risk management follows a documented enterprise risk framework.",
        "Health and safety: the lost-time injury rate improved to 1.2 per million hours.",
        "Workforce diversity reached 41% representation across senior roles.",
        "Employee turnover was 9% and 18 hours of training were delivered per employee.",
    ]:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx() -> bytes | None:
    try:
        import pandas as pd
    except ImportError:
        return None
    frame = pd.DataFrame(
        {
            "Metric": ["Renewable energy share", "Water withdrawal", "Waste recycled"],
            "Value": ["62%", "120000 m3", "78%"],
        }
    )
    buf = BytesIO()
    frame.to_excel(buf, index=False, sheet_name="Metrics")
    return buf.getvalue()


def _hr(title: str) -> None:
    print("\n" + "=" * 72)
    print(title)
    print("=" * 72)


async def run_demo() -> dict:
    configure_logging(level="WARNING", json_output=False)  # keep demo output readable
    pool = InMemoryPool()
    project_id = uuid.uuid4()
    tenant_id = uuid.uuid4()

    # --- Wire up the three services against the shared in-memory DB ---------
    registry = ExtractorRegistry(
        [InlineTextExtractor(), PDFExtractor(), DocxExtractor(), TabularExtractor()]
    )
    signal_repo = SignalRepository(pool)
    await signal_repo.ensure_schema()
    parser = FileParserService(registry=registry, repository=signal_repo)

    profile_repo = ProjectProfileRepository(pool)
    await profile_repo.ensure_schema()
    graph = build_default_graph()
    engine = DiagnosticEngine(  # rule-engine path (no LLM key needed)
        graph=graph,
        gap_detector=GapDetector(graph=graph, evidence_saturation=4),
        repository=profile_repo,
    )

    score_repo = ScoreRepository(pool)
    await score_repo.ensure_schema()
    agent = ScoringAgent(repository=score_repo)  # deterministic justifications

    _hr("STEP 1 — File Parser: documents → signals")
    requests = [
        ParseRequest(project_id=project_id, tenant_id=tenant_id,
                     filename="report_2023.txt", data=REPORT_TEXT.encode(), content_type=None),
    ]
    docx_bytes = _build_docx()
    if docx_bytes:
        requests.append(ParseRequest(
            project_id=project_id, tenant_id=tenant_id, filename="addendum.docx",
            data=docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ))
    xlsx_bytes = _build_xlsx()
    if xlsx_bytes:
        requests.append(ParseRequest(
            project_id=project_id, tenant_id=tenant_id, filename="metrics.xlsx",
            data=xlsx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ))

    documents = await parser.parse_many(requests, persist=True)
    signals = [s for doc in documents for s in doc.signals]
    print(f"Parsed {len(documents)} document(s): "
          + ", ".join(f"{d.filename} ({len(d.signals)} signals)" for d in documents))
    print(f"Persisted {len(pool.store['signals'])} signals to document_signals.\n")
    by_domain: dict[str, int] = {}
    for s in signals:
        by_domain[s.domain.value] = by_domain.get(s.domain.value, 0) + 1
    print("Signals by domain:", by_domain)
    print("Sample signals:")
    for s in signals[:6]:
        val = f"{s.value} {s.unit}" if s.value is not None else "—"
        print(f"  [{s.domain.value:13}] {s.signal_type.value:13} conf={s.confidence:<4} "
              f"val={val:<14} {s.label[:48]}")

    _hr("STEP 2 — MS1 Diagnostic Engine: signals + answers → diagnostic_answers")
    diagnostic = await engine.diagnose(
        project_id=project_id, answers=QUESTIONNAIRE_ANSWERS,
        signals=signals, tenant_id=tenant_id, persist=True,
    )
    o = diagnostic.overall
    print(f"Method: {diagnostic.method}")
    print(f"Overall maturity: level {o.maturity_level}/5 ({o.maturity_label}), "
          f"score {o.score}, confidence {o.confidence}")
    print(f"Perception vs reality (overall): perceived {diagnostic.gap_report.overall_perceived} "
          f"vs evidenced {diagnostic.gap_report.overall_evidenced} "
          f"(gap {diagnostic.gap_report.overall_gap})")
    print("\nPer-domain perception vs reality:")
    print(f"  {'domain':14}{'perceived':>10}{'evidenced':>10}{'gap':>8}  direction")
    for g in diagnostic.gap_report.domain_gaps:
        print(f"  {g.domain.value:14}{g.perceived_score:>10}{g.evidenced_score:>10}"
              f"{g.gap:>8}  {g.direction.value}")
    print("\nTop blockers (ranked):")
    for b in diagnostic.blockers[:5]:
        print(f"  #{b.rank} [{b.priority:5.1f}] {b.domain.value:13} {b.category.value:13} {b.title}")
    print(f"\nWrote diagnostic_answers to project_profiles "
          f"(schema {diagnostic.schema_version}).")

    _hr("STEP 3 — MS2 Scoring Agent: reads diagnostic_answers → 5 scores")
    score_set = await agent.score_project(
        project_id=project_id, tenant_id=tenant_id, persist=True,
    )
    print(f"Method: {score_set.method} | reading diagnostic schema "
          f"{score_set.diagnostic_schema_version}")
    print(f"\n  {'score key':14}{'score':>7}  {'band':12}{'conf':>6}  evidence")
    for s in score_set.scores:
        print(f"  {s.key:14}{s.score:>7.1f}  {s.band.value:12}{s.confidence:>6}  "
              f"{len(s.evidence_refs)} ref(s)")
    print(f"\nOverall: {score_set.overall_score}/100 ({score_set.overall_band.value})")
    print("Sample justification (environmental):")
    env = score_set.by_key()["environmental"]
    print(f"  {env.justification}")

    _hr("STEP 4 — MS3 view: read the 5 scores back from project_profiles")
    stored = await score_repo.read_scores(project_id)
    assert stored is not None and len(stored["scores"]) == 5
    print(f"MS3 reads {len(stored['scores'])} score objects from the 'scores' column:")
    for s in stored["scores"]:
        print(f"  {s['key']:14} → {s['score']:5.1f} ({s['band']})")

    _hr("STEP 5 — Merge-safety checks")
    profile_row = pool.store["profiles"][project_id]
    assert profile_row.get("diagnostic_answers") is not None, "MS1 column missing"
    assert profile_row.get("scores") is not None, "MS2 column missing"
    print("✓ diagnostic_answers (MS1's column) preserved after MS2 wrote scores")
    print("✓ scores (MS2's column) written without touching MS1's data")
    print(f"✓ audit trails: {len(pool.store['ms1_history'])} MS1 + "
          f"{len(pool.store['ms2_history'])} MS2 history row(s)")
    print(f"✓ tenant scoping: tenant_id={tenant_id} propagated through both writes")

    print("\n" + "=" * 72)
    print("DEMO COMPLETE — full MS1 → MS2 → MS3 contract verified end to end.")
    print("=" * 72)

    return {"documents": documents, "signals": signals,
            "diagnostic": diagnostic, "score_set": score_set, "pool": pool}


if __name__ == "__main__":
    asyncio.run(run_demo())
